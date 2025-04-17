import tempfile
import python_docs
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkcalendar import DateEntry
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from io import BytesIO
from pathlib import Path
import os
import atexit

root = tk.Tk()
root.title("Letterer")
root.geometry("600x500")

main_frame = tk.Frame(root)
main_frame.pack(fill="both", expand=True, padx=5, pady=5)

canvas = tk.Canvas(main_frame)
scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
scrollable_frame = tk.Frame(canvas)


scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")



def add_label_entry(row, label_text, widget_type="entry", height=1):
    label = tk.Label(scrollable_frame, text=label_text)
    label.grid(row=row, column=0, sticky="w", padx=5, pady=2)

    if widget_type == "entry":
        entry = tk.Entry(scrollable_frame)
        entry.grid(row=row, column=1, sticky="ew", padx=5, pady=2)
        return entry
    elif widget_type == "text":
        text = tk.Text(scrollable_frame, height=height)
        text.grid(row=row, column=1, sticky="nsew", padx=5, pady=2)
        return text
    elif widget_type == "date":
        date = DateEntry(scrollable_frame, date_pattern='dd.mm.yyyy')
        date.grid(row=row, column=1, sticky="ew", padx=5, pady=2)
        return date


scrollable_frame.grid_columnconfigure(1, weight=1)

org_e = add_label_entry(0, "Название организации:")
from_t = add_label_entry(1, "От кого:", "text", 3)
to_t = add_label_entry(2, "Кому:", "text", 3)
date_d = add_label_entry(3, "Дата:", "date")
purp_e = add_label_entry(4, "Назначение письма:")
l_text_t = add_label_entry(5, "Текст письма:", "text", 6)
pos_e = add_label_entry(6, "Должность отправителя:")
name_e = add_label_entry(7, "Имя отправителя:")

_annexes = "nope"

annex_counter = 0
row_counter = 8
    
annex_list = []


def save_from_buffer():
    context = get_form_data()

    letter_template_path = os.path.join(os.path.dirname(__file__), "Template/letter_template.docx")
    annex_template_path = os.path.join(os.path.dirname(__file__), "Template/annex_template.docx")
    doc = DocxTemplate(letter_template_path)
    annex = DocxTemplate(annex_template_path)
    buffers = []
    
    
    doc.render(context)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    print("Главный документ отрендерен в буфер")
    doc_buffer.seek(0)
    buffers.append(doc_buffer)
    
    for an in annex_list:
        annex_buffer = BytesIO()
        annex.render(get_annex_data(an))
        annex.save(annex_buffer)
        print(an[0].cget("text") + "отрендерен в буфер")
        annex_buffer.seek(0)
        buffers.append(annex_buffer)

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")],
        title="Сохранить документ"
    )

    if file_path:
        merge_docs(buffers, file_path)
        messagebox.showinfo("Успех", "Документ успешно сохранён!")

def preview_from_buffer():
    context = get_form_data()

    letter_template_path = os.path.join(os.path.dirname(__file__), "Template/letter_template.docx")
    annex_template_path = os.path.join(os.path.dirname(__file__), "Template/annex_template.docx")
    doc = DocxTemplate(letter_template_path)
    annex = DocxTemplate(annex_template_path)
    buffers = []
    
    doc.render(context)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    print("Главный документ отрендерен в буфер")
    doc_buffer.seek(0)
    buffers.append(doc_buffer)
    
    for an in annex_list:
        annex_buffer = BytesIO()
        annex.render(get_annex_data(an))
        annex.save(annex_buffer)
        print(an[0].cget("text") + " отрендерен в буфер")
        annex_buffer.seek(0)
        buffers.append(annex_buffer)
        
    merged_doc = Document(buffers[0])
    composer = Composer(merged_doc)
    
    for buf in buffers[1:]:
        doc = Document(buf)
        composer.append(doc)
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        composer.save(tmp_docx.name)
        tmp_docx_path = tmp_docx.name
    
    try:
        os.startfile(tmp_docx_path)
    finally:
        atexit.register(lambda: os.unlink(tmp_docx_path) if os.path.exists(tmp_docx_path) else None)

def merge_docs(buffers, output_path):
    merged_doc = Document(buffers[0])
    composer = Composer(merged_doc)
    
    for file_path in buffers[1:]:
        doc = Document(file_path)
        composer.append(doc)
    
    composer.save(output_path)

def annexes_to_text(list):
    counter = 0
    text = ""
    for an in list:
        counter += 1
        text += an[0].cget("text") + f" на листе {counter}\n"
    return text
    
def get_form_data():
    return {
        "org_name": org_e.get(),
        "from_text": from_t.get("1.0", tk.END).strip(),
        "to_text": to_t.get("1.0", tk.END).strip(),
        "date": date_d.get(),
        "purpose": purp_e.get(),
        "letter_text": l_text_t.get("1.0", tk.END).strip(),
        "sender_position": pos_e.get(),
        "sender_name": name_e.get(),
        "annexes_list": annexes_to_text(annex_list)
    }

def get_annex_data(an):
    return {
        "annex_num": an[0].cget("text"),
        "purpose": purp_e.get(),
        "date": date_d.get(),
        "annex_theme": an[1].get(),
        "annex_text": an[2].get("1.0", tk.END),
    }
 
def destroy_annex(name):
    global annex_counter
    for an in annex_list:
        if an[0].cget("text") == name:
            print(f"Удаление {name}")
            for el in an:
                print("Уничтожен виджет")
                el.destroy()
            annex_list.remove(an)
            print("Элемент удалён")
            break
    
    counter = 1
    for an in annex_list:
        an[0].config(text=f"Приложение {counter}")
        an[5].config(text=f"Удалить приложение {counter}")
        counter += 1
        
    annex_counter -= 1
        
def generate_annex(an_btn, btn):
    global root
    global annex_counter
    global row_counter
    global annex_list
    annex_counter += 1
    this_annex = []
    
    row_counter += 1
    an_lab = tk.Label(scrollable_frame, text="Приложение " + str(annex_counter))
    an_lab.grid(row=row_counter, column=0, columnspan=2, sticky="w")
    this_annex.append(an_lab)
    
    row_counter += 1
    purp_label = tk.Label(scrollable_frame, text="Назначение приложения")
    purp_label.grid(row=row_counter, column=0, sticky="w", padx=5, pady=2)
    an_purp = tk.Entry(scrollable_frame)
    an_purp.grid(row=row_counter, column=1, sticky="ew", padx=5, pady=2)
    this_annex.append(an_purp)
    
    row_counter += 1
    text_label = tk.Label(scrollable_frame, text="Текст приложения")
    text_label.grid(row=row_counter, column=0, sticky="w", padx=5, pady=2)
    an_text = tk.Text(scrollable_frame, height=10)
    an_text.grid(row=row_counter, column=1, sticky="nsew", padx=5, pady=2)
    this_annex.append(an_text)
    
    row_counter += 1
    name = this_annex[0].cget("text")
    del_btn = tk.Button(scrollable_frame, text=f"Удалить приложение {annex_counter}", command=lambda: destroy_annex(this_annex[0].cget("text")))
    del_btn.grid(row=row_counter, column=0, pady=10)
    
    this_annex.append(purp_label)
    this_annex.append(text_label)
    this_annex.append(del_btn)
    
    row_counter += 1
    an_btn.grid(row=row_counter, column=0, columnspan=2, pady=10)
    
    row_counter += 1
    prev_btn.grid(row=row_counter, column=0, pady=10)
    btn.grid(row=row_counter, column=1, pady=10)
    
    annex_list.append(this_annex)



btn = tk.Button(scrollable_frame, text="Сохранить документ", command=save_from_buffer)
btn.grid(row=9, column=1, pady=10)

prev_btn = tk.Button(scrollable_frame, text="Предпросмотр", command=preview_from_buffer)
prev_btn.grid(row=9, column=0, pady=10)

an_btn = tk.Button(scrollable_frame, text="Создать приложение", command=lambda: generate_annex(an_btn, btn))
an_btn.grid(row=8, column=0, columnspan=2, pady=10)


root.mainloop()